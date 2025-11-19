#!/usr/bin/env python3
# -*- coding: utf-8 -*-

"""
Test script for DOCX hyperlink functionality
"""

from docx import Document
from docx.shared import RGBColor
from docx.oxml.ns import nsdecls, qn
from docx.oxml import parse_xml
import docx

def test_docx_hyperlinks():
    """Test creating DOCX with hyperlinks"""

    # Create a new document
    doc = Document()

    # Set document title
    doc.add_heading('测试DOCX超链接功能', 0)

    # Test 1: Simple text with styled link (fallback)
    para1 = doc.add_paragraph()
    para1.add_run("这是第一个测试链接：")
    link_run = para1.add_run("https://www.example.com")
    link_run.font.color.rgb = RGBColor(0, 0, 255)  # Blue
    link_run.font.underline = True  # Underlined

    # Test 2: Try to create actual hyperlink
    para2 = doc.add_paragraph()
    para2.add_run("这是第二个测试链接：")

    try:
        # Create hyperlink relationship
        r_id = doc.part.relate_to("https://www.example.com", docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

        # Create hyperlink element
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
        hyperlink.set(docx.oxml.shared.qn('r:id'), r_id)

        # Create run element with styling
        run = docx.oxml.shared.OxmlElement('w:r')

        # Add run properties (styling)
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        color = docx.oxml.shared.OxmlElement('w:color')
        color.set(docx.oxml.shared.qn('w:val'), '0000FF')  # Blue
        underline = docx.oxml.shared.OxmlElement('w:u')
        underline.set(docx.oxml.shared.qn('w:val'), 'single')  # Underlined
        rPr.append(color)
        rPr.append(underline)
        run.append(rPr)

        # Add text
        text_elem = docx.oxml.shared.OxmlElement('w:t')
        text_elem.text = "https://www.example.com"
        run.append(text_elem)

        hyperlink.append(run)
        para2._element.append(hyperlink)

        print("✅ 成功创建超链接元素")

    except Exception as e:
        print("❌ 创建超链接失败: {}".format(e))
        # Fallback to styled text
        fallback_run = para2.add_run("https://www.example.com (备用样式)")
        fallback_run.font.color.rgb = RGBColor(0, 0, 255)
        fallback_run.font.underline = True

    # Save the document
    doc.save('test_docx_links.docx')
    print("✅ DOCX文件已保存为 test_docx_links.docx")

if __name__ == "__main__":
    test_docx_hyperlinks()
